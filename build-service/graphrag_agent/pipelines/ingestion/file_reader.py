import codecs
import os
from typing import List, Tuple, Dict, Optional
import PyPDF2
from docx import Document
import csv
import json
import yaml
from yaml import CLoader as Loader

from graphrag_agent.config.settings import FILES_DIR


class FileReader:
    """
    File reader that supports multiple file formats:
    - TXT (text files)
    - PDF (PDF documents)
    - MD (Markdown files)
    - DOCX (Word documents)
    - DOC (legacy Word documents)
    - CSV (CSV files)
    - JSON (JSON files)
    - YAML/YML (YAML files)
    """

    def __init__(self, directory_path: str):
        """
        Initialize the file reader.
        
        Args:
            directory_path: File directory path
        """
        self.directory_path = directory_path
        
    def read_files(self, file_extensions: Optional[List[str]] = None, recursive: bool = True) -> List[Tuple[str, str]]:
        """
        Read files with specified extensions.
        
        Args:
            file_extensions: List of file extensions (e.g., ['.txt', '.pdf']); if None, read all supported formats
            recursive: Whether to read subdirectories recursively (default True)
            
        Returns:
            List[Tuple[str, str]]: List of (filename, content) tuples
        """
        supported_extensions = {
            '.txt': self._read_txt,
            '.pdf': self._read_pdf,
            '.md': self._read_markdown,
            '.docx': self._read_docx,
            '.doc': self._read_doc,
            '.csv': self._read_csv,
            '.json': self._read_json,
            '.yaml': self._read_yaml,
            '.yml': self._read_yaml,
        }
        
        # If no extensions are specified, use all supported extensions
        if file_extensions is None:
            file_extensions = list(supported_extensions.keys())
            
        results = []
        try:
            if recursive:
                # Recursively read all files
                results = self._read_files_recursive(self.directory_path, file_extensions, supported_extensions)
                print(f"Recursive read complete; total files read: {len(results)}")
            else:
                # Read files only in the current directory
                all_filenames = os.listdir(self.directory_path)
                print(f"Current directory contains {len(all_filenames)} files")
                
                results = self._process_files_in_dir(self.directory_path, all_filenames, file_extensions, supported_extensions)
                print(f"Total files read: {len(results)}")
        except Exception as e:
            print(f"Error listing files in directory {self.directory_path}: {str(e)}")
            
        return results
    
    def _read_files_recursive(self, root_dir: str, file_extensions: List[str], supported_extensions: Dict) -> List[Tuple[str, str]]:
        """
        Recursively read files in a directory and its subdirectories.
        
        Args:
            root_dir: Current directory path
            file_extensions: List of file extensions to process
            supported_extensions: Supported extensions and their handlers
            
        Returns:
            List[Tuple[str, str]]: List of (filename, content) tuples
        """
        results = []
        
        try:
            # Traverse directory contents
            for item in os.listdir(root_dir):
                item_path = os.path.join(root_dir, item)
                
                # If it's a directory, recurse into it
                if os.path.isdir(item_path):
                    print(f"Recursing into subdirectory: {item_path}")
                    sub_results = self._read_files_recursive(item_path, file_extensions, supported_extensions)
                    results.extend(sub_results)
                
                # If it's a file, process it
                elif os.path.isfile(item_path):
                    file_ext = os.path.splitext(item)[1].lower()
                    
                    if file_ext in file_extensions:
                        # Get the path relative to the root directory
                        rel_path = os.path.relpath(item_path, self.directory_path)
                        
                        print(f"Processing file: {rel_path} (type: {file_ext})")
                        
                        # Use the corresponding reader for the file type
                        if file_ext in supported_extensions:
                            try:
                                content = supported_extensions[file_ext](item_path)
                                # Store the relative path (not just filename) to distinguish files with the same name
                                results.append((rel_path, content))
                                print(f"Successfully read file: {rel_path}, content length: {len(content)}")
                            except Exception as e:
                                print(f"Error reading file {rel_path}: {str(e)}")
        except Exception as e:
            print(f"Error listing files in directory {root_dir}: {str(e)}")
            
        return results
    
    def _process_files_in_dir(self, directory: str, filenames: List[str], file_extensions: List[str], 
                              supported_extensions: Dict) -> List[Tuple[str, str]]:
        """
        Process files in a specific directory (non-recursive).
        
        Args:
            directory: Directory path
            filenames: List of filenames
            file_extensions: List of file extensions to process
            supported_extensions: Supported extensions and their handlers
            
        Returns:
            List[Tuple[str, str]]: List of (filename, content) tuples
        """
        results = []
        
        for filename in filenames:
            file_ext = os.path.splitext(filename)[1].lower()
            
            if file_ext in file_extensions:
                file_path = os.path.join(directory, filename)
                print(f"Processing file: {filename} (type: {file_ext})")
                
                # Use the corresponding reader for the file type
                if file_ext in supported_extensions:
                    try:
                        content = supported_extensions[file_ext](file_path)
                        results.append((filename, content))
                        print(f"Successfully read file: {filename}, content length: {len(content)}")
                    except Exception as e:
                        print(f"Error reading file {filename}: {str(e)}")
        
        return results
    
    def _read_txt(self, file_path: str) -> str:
        """Read a TXT file."""
        try:
            with codecs.open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                content = file.read()
            return content
        except Exception as e:
            print(f"Failed to read TXT file {os.path.basename(file_path)}: {str(e)}")
            # Try another encoding
            try:
                with open(file_path, 'rb') as f:
                    raw_data = f.read(10240)  # Read first 10KB
                    try:
                        import chardet
                        result = chardet.detect(raw_data)
                        encoding = result['encoding'] if result['encoding'] else 'gbk'
                    except:
                        encoding = 'gbk'  # Default to gbk if chardet is unavailable
                        
                with codecs.open(file_path, 'r', encoding=encoding, errors='replace') as file:
                    content = file.read()
                return content
            except Exception as e2:
                print(f"Failed to read with alternate encoding: {str(e2)}")
                return f"[Unable to read file content: {str(e)}]"
            
    def _read_pdf(self, file_path: str) -> str:
        """Read a PDF file."""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    try:
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text() or ""
                        text += page_text + "\n\n"
                    except Exception as e:
                        print(f"Failed to read page {page_num+1} of PDF {os.path.basename(file_path)}: {str(e)}")
                        text += f"[Page {page_num+1} could not be read]\n\n"
            return text
        except Exception as e:
            print(f"Failed to read PDF file {os.path.basename(file_path)}: {str(e)}")
            return f"[Unable to read PDF content: {str(e)}]"
    
    def _read_markdown(self, file_path: str) -> str:
        """Read a Markdown file."""
        try:
            with codecs.open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                md_content = file.read()
                return md_content
        except Exception as e:
            print(f"Failed to read Markdown file {os.path.basename(file_path)}: {str(e)}")
            return f"[Unable to read Markdown content: {str(e)}]"
    
    def _read_docx(self, file_path: str) -> str:
        """Read a Word document (.docx)."""
        try:
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return '\n'.join(full_text)
        except Exception as e:
            print(f"Failed to read Word document (.docx) {os.path.basename(file_path)}: {str(e)}")
            return f"[Unable to read Word document content: {str(e)}]"
            
    def _read_doc(self, file_path: str) -> str:
        """
        Read a legacy Word document (.doc).
        First try a Windows-specific method, then fall back to cross-platform methods if needed.
        """
        content = ""
        
        # Method 1: try win32com (Windows only)
        try:
            import win32com.client
            
            print(f"Trying win32com to read .doc file: {os.path.basename(file_path)}")
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            doc_abs_path = os.path.abspath(file_path)
            doc = word.Documents.Open(doc_abs_path)
            content = doc.Content.Text
            doc.Close()
            word.Quit()
            
            if content and content.strip():
                print("Successfully read .doc file with win32com")
                return content
        except ImportError:
            print("win32com is unavailable; this is not a Windows system")
        except Exception as e:
            print(f"win32com failed to read .doc: {str(e)}")
        
        # Method 2: try textract (cross-platform)
        try:
            import textract
            print(f"Trying textract to read .doc file: {os.path.basename(file_path)}")
            content = textract.process(file_path).decode('utf-8')
            
            if content and content.strip():
                print("Successfully read .doc file with textract")
                return content
        except ImportError:
            print("textract is unavailable; install with: pip install textract")
        except Exception as e:
            print(f"textract failed to read .doc: {str(e)}")
        
        # Method 3: try python-docx (not fully compatible with .doc, but may partially work)
        try:
            from docx import Document
            print(f"Trying python-docx to read .doc file: {os.path.basename(file_path)}")
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            content = '\n'.join(full_text)
            
            if content and content.strip():
                print("Successfully partially read .doc file with python-docx")
                return content
        except ImportError:
            print("python-docx is unavailable; install with: pip install python-docx")
        except Exception as e:
            print(f"python-docx failed to read .doc: {str(e)}")
        
        # If all methods fail, return a warning message
        warning_msg = f"[Warning: Failed to read .doc file {os.path.basename(file_path)}. Please install dependencies or convert to .docx]"
        print(warning_msg)
        return warning_msg

    
    def _read_csv(self, file_path: str) -> str:
        """
        Read a CSV file and convert it to text.
        
        Note: This converts CSV to plain text and does not preserve structured data.
        """
        try:
            text = []
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                csv_reader = csv.reader(file)
                for row in csv_reader:
                    text.append(','.join(row))
            return '\n'.join(text)
        except Exception as e:
            print(f"Failed to read CSV file {os.path.basename(file_path)}: {str(e)}")
            # Try another encoding
            try:
                with open(file_path, 'rb') as f:
                    try:
                        import chardet
                        raw_data = f.read(10240)
                        result = chardet.detect(raw_data)
                        encoding = result['encoding'] if result['encoding'] else 'gbk'
                    except:
                        encoding = 'gbk'  # Default to gbk if chardet is unavailable
                        
                text = []
                with open(file_path, 'r', encoding=encoding, errors='replace') as file:
                    csv_reader = csv.reader(file)
                    for row in csv_reader:
                        text.append(','.join(row))
                return '\n'.join(text)
            except Exception as e2:
                print(f"Failed to read CSV with alternate encoding: {str(e2)}")
                return f"[Unable to read CSV content: {str(e)}]"
    
    def read_csv_as_dicts(self, file_path: str) -> List[Dict]:
        """
        Read a CSV file and return a list of dicts.
        
        Returns:
            List[Dict]: CSV rows as dictionaries
        """
        try:
            results = []
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                csv_reader = csv.DictReader(file)
                for row in csv_reader:
                    results.append(dict(row))
            return results
        except Exception as e:
            print(f"Error reading CSV as dicts: {str(e)}")
            return []
    
    def _read_json(self, file_path: str) -> str:
        """Read a JSON file and return text."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                # Load as an object, then serialize to a formatted string for readability
                data = json.load(file)
                return json.dumps(data, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Failed to read JSON file {os.path.basename(file_path)}: {str(e)}")
            return f"[Unable to read JSON content: {str(e)}]"
    
    def read_json_as_dict(self, file_path: str) -> Dict:
        """
        Read a JSON file and return a dict/list object.
        
        Returns:
            Dict/List: JSON data object
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                return json.load(file)
        except Exception as e:
            print(f"Error reading JSON as dict: {str(e)}")
            return {}
    
    def _read_yaml(self, file_path: str) -> str:
        """Read a YAML file and return text."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                data = yaml.load(file, Loader=Loader)
                # Convert to YAML string for a more readable format
                return yaml.dump(data, allow_unicode=True, default_flow_style=False)
        except Exception as e:
            print(f"Failed to read YAML file {os.path.basename(file_path)}: {str(e)}")
            return f"[Unable to read YAML content: {str(e)}]"
    
    def read_yaml_as_dict(self, file_path: str) -> Dict:
        """
        Read a YAML file and return a dict object.
        
        Returns:
            Dict: YAML data object
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                return yaml.load(file, Loader=Loader)
        except Exception as e:
            print(f"Error reading YAML as dict: {str(e)}")
            return {}
    
    def read_txt_files(self) -> List[Tuple[str, str]]:
        """Read all txt files."""
        return self.read_files(['.txt'])
    
    def list_all_files(self, recursive: bool = True) -> List[str]:
        """
        List all files in a directory.
        
        Args:
            recursive: Whether to list files in subdirectories (default True)
            
        Returns:
            List[str]: File paths relative to the root directory
        """
        files = []
        
        try:
            if recursive:
                # Recursively traverse all subdirectories
                for root, _, filenames in os.walk(self.directory_path):
                    for filename in filenames:
                        # Get the path relative to the root directory
                        rel_path = os.path.relpath(os.path.join(root, filename), self.directory_path)
                        files.append(rel_path)
            else:
                # List files only in the current directory
                files = os.listdir(self.directory_path)
        except Exception as e:
            print(f"Error listing directory files: {str(e)}")
            
        return files


# Test code
if __name__ == '__main__':
    print(f"FILES_DIR: {FILES_DIR}")
    reader = FileReader(FILES_DIR)
    
    # List all files in the directory
    all_filenames = reader.list_all_files()
    print(f"Directory contains {len(all_filenames)} files:")
    for filename in all_filenames:
        print(f"  {filename}")
    
    # Test reading all supported files
    all_files = reader.read_files()
    print(f"Successfully read {len(all_files)} files")
    
    # Show the number of files by type
    file_types = {}
    for file_name, _ in all_files:
        ext = os.path.splitext(file_name)[1].lower()
        file_types[ext] = file_types.get(ext, 0) + 1
    
    print("Files by type:")
    for ext, count in file_types.items():
        print(f"  {ext}: {count}")